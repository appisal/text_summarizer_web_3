import streamlit as st
from io import BytesIO
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import seaborn as sns
import urllib.parse  
from reportlab.pdfgen import canvas
from docx import Document
import pdfplumber
from gtts import gTTS
import qrcode
from PIL import Image
from io import BytesIO
import time
st.markdown("""
    <style>
        /* Apply background to the whole page */
        .stApp {
            background: linear-gradient(to bottom, #f8f9fa, #e3e7ed);
        }

        /* Style for text areas and file upload boxes */
        .stTextArea, .stFileUploader {
            background-color: rgba(255, 255, 255, 0.8);
            border-radius: 10px;
            padding: 10px;
            box-shadow: 2px 2px 10px rgba(0, 0, 0, 0.1);
        }

        /* Style the headers for better visibility */
        h1 {
            color: #2c3e50;
            text-align: center;
        }
        
        /* Typing animation for summary */
        @keyframes typing {
            from { width: 0; }
            to { width: 100%; }
        }
        .summary-text {
            overflow: hidden;
            white-space: nowrap;
            border-right: 2px solid black;
            animation: typing 2s steps(40, end) forwards;
        }
        
        /* Fade-in effect for QR code */
        @keyframes fadeIn {
            from { opacity: 0; }
            to { opacity: 1; }
        }
        .qr-container img {
            animation: fadeIn 1s ease-in-out;
        }
        
        /* Button effects */
        .stButton>button {
            transition: transform 0.2s ease-in-out;
        }
        .stButton>button:hover {
            transform: scale(1.1);
        }
    </style>
""", unsafe_allow_html=True)




# Session state for history
if "summary_history" not in st.session_state:
    st.session_state.summary_history = []

# PDF & DOCX Upload Support
def extract_text_from_pdf(uploaded_file):
    with pdfplumber.open(uploaded_file) as pdf:
        return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])

def extract_text_from_docx(uploaded_file):
    doc = Document(uploaded_file)
    return "\n".join([para.text for para in doc.paragraphs])

# Function to download summary as PDF
def download_pdf(summary):
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(100, 750, "Summary Report")
    pdf.drawString(100, 730, summary)
    pdf.save()
    buffer.seek(0)
    return buffer

# Function to download summary as Word
def download_word(summary):
    doc = Document()
    doc.add_paragraph(summary)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to download summary as audio
def download_audio(summary):
    tts = gTTS(summary, lang="en")
    buffer = BytesIO()
    tts.write_to_fp(buffer)
    buffer.seek(0)
    return buffer

# Function to generate sharing links

def generate_share_links(summary):
    encoded_summary = urllib.parse.quote(summary)
    return {
        "WhatsApp": f"https://wa.me/?text={encoded_summary}",
        "Twitter": f"https://twitter.com/intent/tweet?text={encoded_summary}",
        "Email": f"mailto:?subject=Summary&body={encoded_summary}",
        "LinkedIn": f"https://www.linkedin.com/sharing/share-offsite/?url={encoded_summary}",
        "Facebook": f"https://www.facebook.com/sharer/sharer.php?u={encoded_summary}",
        "Telegram": f"https://t.me/share/url?url={encoded_summary}"
    }

# Function to generate a QR code with smaller size
def generate_qr_code(summary):
    qr = qrcode.QRCode(
        version=1,  # Controls the complexity (1 is smallest)
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=3,  # Reduce this for a smaller QR code
        border=2     # Controls the white border (default is 4)
    )
    qr.add_data(summary)
    qr.make(fit=True)

    img = qr.make_image(fill="black", back_color="white")

    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer

# Function to create share buttons with icons
def create_share_buttons(summary):
    share_links = generate_share_links(summary)

    st.markdown("""
    <style>
        .share-btn-container {
            display: flex;
            justify-content: center;
            gap: 15px;
            margin-top: 20px;
        }
        .share-btn img {
            width: 40px;
            height: 40px;
            transition: transform 0.3s ease-in-out;
            cursor: pointer;
        }
        .share-btn img:hover {
            transform: scale(1.2);
        }
    </style>
    """, unsafe_allow_html=True)

    # Share buttons with clickable icons
    share_html = f"""
    <div class="share-btn-container">
        <a href="{share_links['WhatsApp']}" target="_blank" class="share-btn">
            <img src="https://upload.wikimedia.org/wikipedia/commons/6/6b/WhatsApp.svg" alt="WhatsApp">
        </a>
        <a href="{share_links['Twitter']}" target="_blank" class="share-btn">
            <img src="https://upload.wikimedia.org/wikipedia/en/6/60/Twitter_Logo_as_of_2021.svg" alt="Twitter">
        </a>
        <a href="{share_links['Email']}" target="_blank" class="share-btn">
            <img src="https://upload.wikimedia.org/wikipedia/commons/4/4e/Mail_%28iOS%29.svg" alt="Email">
        </a>
        <a href="{share_links['LinkedIn']}" target="_blank" class="share-btn">
            <img src="https://upload.wikimedia.org/wikipedia/commons/c/ca/LinkedIn_logo_initials.png" alt="LinkedIn">
        </a>
        <a href="{share_links['Facebook']}" target="_blank" class="share-btn">
            <img src="https://upload.wikimedia.org/wikipedia/commons/5/51/Facebook_f_logo_%282019%29.svg" alt="Facebook">
        </a>
        <a href="{share_links['Telegram']}" target="_blank" class="share-btn">
            <img src="https://upload.wikimedia.org/wikipedia/commons/8/82/Telegram_logo.svg" alt="Telegram">
        </a>
    </div>
    """

    st.markdown(share_html, unsafe_allow_html=True)

    # Display QR code
    st.markdown("<h4 style='text-align: center;'>📲 Scan QR Code to Share</h4>", unsafe_allow_html=True)
    qr_buffer = generate_qr_code(summary)
    st.image(qr_buffer, caption="QR Code for Sharing", use_container_width=False)
   
   
  

# UI Setup
st.markdown("""
    <div style="text-align: center;">
        <img src="https://i.imgur.com/55eTSwu.png" alt="QuickText Logo" style="width: 120px; margin-bottom: -10px;">
        <h1 style="margin-top: 5px;">🚀 QuickText - Text Processor</h1>
    </div>
""", unsafe_allow_html=True)

st.sidebar.title("⚡ Features")
option = st.sidebar.radio("Choose an option:", ["Single File", "Bulk File Processing", "History"])

if option == "Single File":
    st.markdown("<h3>📂 Upload a file or paste text.</h3>", unsafe_allow_html=True)
    uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx"])
    text = ""

    if uploaded_file:
        if uploaded_file.type == "application/pdf":
            text = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_text_from_docx(uploaded_file)
    else:
        text = st.text_area("✍️ Paste your text here:", height=200)

    if text.strip():
        col1, col2 = st.columns(2)
        with col1:
            min_length = st.slider("📏 Min length:", 10, 100, 50)
        with col2:
            max_length = st.slider("📏 Max length:", 50, 500, 200)

        # ✅ NEW: Live Summary Preview with Animation & Highlight
        preview_text = text[:max_length]
        st.markdown("""
        <style>
            .summary-box {
                padding: 10px;
                border-radius: 8px;
                background: #f0f8ff;
                font-weight: bold;
                animation: fadeIn 0.5s ease-in-out;
            }
            @keyframes fadeIn {
                0% { opacity: 0; transform: translateY(-10px); }
                100% { opacity: 1; transform: translateY(0); }
            }
        </style>
        """, unsafe_allow_html=True)

        st.markdown("<h3>👀 Live Summary Preview:</h3>", unsafe_allow_html=True)
        st.markdown(f'<div class="summary-box">{preview_text}</div>', unsafe_allow_html=True)

        if st.button("📑 Process", use_container_width=True):
            processed_text = text[:max_length]
            st.markdown("<h3>📌 Processed Text:</h3>", unsafe_allow_html=True)
            st.success(processed_text)

            # Download Buttons
            col1, col2, col3 = st.columns(3)
            with col1:
                st.download_button("📄 PDF", download_pdf(processed_text), file_name="processed_text.pdf", mime="application/pdf")
            with col2:
                st.download_button("📝 Word", download_word(processed_text), file_name="processed_text.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with col3:
                st.download_button("🔊 Audio", download_audio(processed_text), file_name="processed_text.mp3", mime="audio/mp3")

            # ✅ NEW: Share Buttons
            st.markdown("<h3 style='text-align: center;'>📢 Share</h3>", unsafe_allow_html=True)
            create_share_buttons(processed_text)

elif option == "Bulk File Processing":
    uploaded_files = st.file_uploader("Upload multiple files", type=["pdf", "docx"], accept_multiple_files=True)
    if uploaded_files:
        for file in uploaded_files:
            text = extract_text_from_pdf(file) if file.type == "application/pdf" else extract_text_from_docx(file)
            processed_text = text[:200]
            st.markdown(f"### 📜 Processed Text for {file.name}")
            st.success(processed_text)

elif option == "History":
    st.subheader("📜 History")
    for i, text in enumerate(reversed(st.session_state.summary_history)):
        with st.expander(f"📄 Entry {len(st.session_state.summary_history) - i}"):
            st.write(text)
            create_share_buttons(text)

st.markdown("<hr><p style='text-align: center;'>🔗 QuickText - Text Processor</p>", unsafe_allow_html=True)
